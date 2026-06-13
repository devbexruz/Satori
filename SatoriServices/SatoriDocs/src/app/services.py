import io
import re
from typing import Optional, List
from uuid import UUID
from sqlalchemy.orm import Session
import docx
from docx.shared import Inches

from app.config import settings
from app.models import Document, Module, Section
from app.schemas import DocumentOutline, ModuleOutline, SectionOutline

from langchain_google_genai import ChatGoogleGenerativeAI
from langchain_core.messages import SystemMessage, HumanMessage


# --- EXPORT TO MARKDOWN ---

def export_document_to_markdown(document: Document) -> str:
    """
    Renders the entire Document into a single cohesive Markdown string.
    """
    lines = []
    lines.append(f"# {document.title}")
    if document.description:
        lines.append(document.description)
    lines.append("")  # Blank line after title/desc

    # Sort modules by their order
    sorted_modules = sorted(document.modules, key=lambda m: m.order)
    
    for module in sorted_modules:
        lines.append(f"## {module.title}")
        if module.description:
            lines.append(module.description)
        lines.append("")

        # Fetch top level sections of this module (parent_id is None)
        top_sections = [s for s in module.sections if s.parent_id is None]
        top_sections.sort(key=lambda s: s.order)

        def render_section(section: Section):
            # level is 1 to 5. We prefix with (level + 2) hash marks
            # e.g., Section Level 1 becomes H3: '### SectionTitle'
            heading_prefix = "#" * (section.level + 2)
            lines.append(f"{heading_prefix} {section.title}")
            if section.content:
                lines.append(section.content)
            lines.append("")

            # Render children nested sections
            sorted_children = sorted(section.children, key=lambda c: c.order)
            for child in sorted_children:
                render_section(child)

        for section in top_sections:
            render_section(section)

    return "\n".join(lines)


# --- EXPORT TO JSON ---

def export_document_to_json(document: Document) -> dict:
    """
    Serializes the entire Document structure including its Modules and nested Sections into JSON.
    """
    def get_section_dict(section: Section) -> dict:
        return {
            "id": str(section.id),
            "parent_id": str(section.parent_id) if section.parent_id else None,
            "title": section.title,
            "level": section.level,
            "content": section.content,
            "order": section.order,
            "children": [get_section_dict(c) for c in sorted(section.children, key=lambda x: x.order)]
        }

    return {
        "id": str(document.id),
        "title": document.title,
        "description": document.description,
        "created_at": document.created_at.isoformat(),
        "updated_at": document.updated_at.isoformat(),
        "modules": [
            {
                "id": str(module.id),
                "title": module.title,
                "description": module.description,
                "order": module.order,
                "sections": [
                    get_section_dict(s)
                    for s in sorted(module.sections, key=lambda x: x.order)
                    if s.parent_id is None
                ]
            }
            for module in sorted(document.modules, key=lambda m: m.order)
        ]
    }


# --- EXPORT TO WORD (.docx) ---

def add_formatted_text(paragraph, text: str):
    """
    Parses basic inline Markdown (bold **, italic *, inline code `)
    and applies native formatting to docx runs.
    """
    if not text:
        return
    # Regex to split text by formatting tokens: **bold**, *italic*, `code`
    parts = re.split(r"(\*\*.*?\*\*|\*.*?\*|`.*?`)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            # Bold
            paragraph.add_run(part[2:-2]).bold = True
        elif part.startswith("*") and part.endswith("*"):
            # Italic
            paragraph.add_run(part[1:-1]).italic = True
        elif part.startswith("`") and part.endswith("`"):
            # Inline code
            run = paragraph.add_run(part[1:-1])
            run.font.name = "Courier New"
        else:
            paragraph.add_run(part)


def parse_markdown_content(doc: docx.Document, md_text: str):
    """
    Parses a Markdown block text and appends formatted elements (lists, code blocks, quotes)
    as native elements in a python-docx document.
    """
    if not md_text:
        return

    lines = md_text.splitlines()
    in_code_block = False
    code_lines = []
    
    i = 0
    while i < len(lines):
        line = lines[i]
        
        # Code block check
        if line.strip().startswith("```"):
            if in_code_block:
                in_code_block = False
                code_text = "\n".join(code_lines)
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Inches(0.5)
                run = p.add_run(code_text)
                run.font.name = "Courier New"
                code_lines = []
            else:
                in_code_block = True
            i += 1
            continue
            
        if in_code_block:
            code_lines.append(line)
            i += 1
            continue
            
        stripped = line.strip()
        if not stripped:
            i += 1
            continue
            
        # Heading checks inside sections (e.g. # title, ## title)
        if stripped.startswith("#"):
            h_match = re.match(r"^(#{1,6})\s+(.*)$", stripped)
            if h_match:
                level = len(h_match.group(1))
                text = h_match.group(2)
                # Word supports levels 1-9
                doc.add_heading(text, level=min(level, 9))
                i += 1
                continue
                
        # Blockquote check
        if stripped.startswith(">"):
            quote_text = stripped[1:].strip()
            # Group multi-line blockquotes together
            while i + 1 < len(lines) and lines[i+1].strip().startswith(">"):
                i += 1
                quote_text += "\n" + lines[i].strip()[1:].strip()
            p = doc.add_paragraph(style="Quote")
            add_formatted_text(p, quote_text)
            i += 1
            continue
            
        # Bullet list check (starts with -, *, +)
        if stripped.startswith("- ") or stripped.startswith("* ") or stripped.startswith("+ "):
            bullet_text = stripped[2:].strip()
            p = doc.add_paragraph(style="List Bullet")
            add_formatted_text(p, bullet_text)
            i += 1
            continue
            
        # Numbered list check (starts with digits e.g. "1. ")
        num_match = re.match(r"^(\d+)\.\s+(.*)$", stripped)
        if num_match:
            item_text = num_match.group(2).strip()
            p = doc.add_paragraph(style="List Number")
            add_formatted_text(p, item_text)
            i += 1
            continue
            
        # Regular text paragraph - can wrap multiple consecutive text lines
        paragraph_lines = [stripped]
        while (i + 1 < len(lines) 
               and lines[i+1].strip() 
               and not any(lines[i+1].strip().startswith(x) for x in ["#", ">", "- ", "* ", "+ "])
               and not re.match(r"^(\d+)\.\s+(.*)$", lines[i+1].strip())
               and not lines[i+1].strip().startswith("```")):
            i += 1
            paragraph_lines.append(lines[i].strip())
            
        p = doc.add_paragraph()
        add_formatted_text(p, " ".join(paragraph_lines))
        i += 1


def export_document_to_docx(document: Document) -> io.BytesIO:
    """
    Generates a Microsoft Word (.docx) file representing the document,
    using native styles and properly parsing markdown structures.
    """
    doc = docx.Document()
    
    # Document title
    doc.add_heading(document.title, level=0)
    if document.description:
        p = doc.add_paragraph()
        add_formatted_text(p, document.description)
        
    # Order modules by order
    sorted_modules = sorted(document.modules, key=lambda m: m.order)
    
    for module in sorted_modules:
        # Module title
        doc.add_heading(module.title, level=1)
        if module.description:
            p = doc.add_paragraph()
            add_formatted_text(p, module.description)
            
        # Top-level sections of module
        top_sections = [s for s in module.sections if s.parent_id is None]
        top_sections.sort(key=lambda s: s.order)
        
        def render_section_docx(section: Section):
            # Section Title. Level is 1-5, so heading is level + 1 (Heading 2, 3, etc.)
            doc.add_heading(section.title, level=min(section.level + 1, 9))
            
            if section.content:
                parse_markdown_content(doc, section.content)
                
            # Render nested sections
            sorted_children = sorted(section.children, key=lambda c: c.order)
            for child in sorted_children:
                render_section_docx(child)
                
        for section in top_sections:
            render_section_docx(section)
            
    file_stream = io.BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)
    return file_stream


# --- AI GENERATION WITH LANGCHAIN & GEMINI ---

def check_gemini_config():
    """Raises ValueError if Gemini API key is missing."""
    if not settings.GEMINI_API_KEY:
        raise ValueError("GEMINI_API_KEY environment variable is not configured. AI capabilities are disabled.")


def generate_document_outline(topic: str, description: Optional[str] = None) -> DocumentOutline:
    """
    Uses Gemini structured output via LangChain to generate modules and sections structure for a document.
    """
    check_gemini_config()
    
    llm = ChatGoogleGenerativeAI(
        model=settings.GEMINI_MODEL,
        google_api_key=settings.GEMINI_API_KEY,
        temperature=0.7
    )
    
    # Enable structured output based on the DocumentOutline Pydantic schema
    structured_llm = llm.with_structured_output(DocumentOutline)
    
    prompt = f"Create a comprehensive document outline for the topic: '{topic}'."
    if description:
        prompt += f"\nDescription/Context: {description}"
        
    outline = structured_llm.invoke(prompt)
    return outline


def generate_section_content(
    section_title: str,
    level: int,
    module_title: str,
    context: Optional[str] = None,
    prompt_instruction: Optional[str] = None
) -> str:
    """
    Generates rich markdown content for a section based on its heading, level, module and overall context.
    """
    check_gemini_config()
    
    llm = ChatGoogleGenerativeAI(
        model=settings.GEMINI_MODEL,
        google_api_key=settings.GEMINI_API_KEY,
        temperature=0.7
    )
    
    system_msg = SystemMessage(content=(
        "You are an expert technical writer and document generator. "
        "Your task is to write detailed, high-quality content for a specific document section. "
        "Write in rich, valid Markdown format. Include list items (both ordered and unordered), bold/italic texts, links, blockquotes, "
        "code blocks, or small tables where appropriate to make the section rich and informative. "
        "Do not include the heading itself at the start of your response, only output the content under the heading."
    ))
    
    user_content = f"Write the section content for section heading: '{section_title}' (heading level H{level}).\n"
    user_content += f"This section is part of the module: '{module_title}'.\n"
    if context:
        user_content += f"Document context:\n{context}\n"
    if prompt_instruction:
        user_content += f"Additional generation instruction: {prompt_instruction}\n"
        
    user_msg = HumanMessage(content=user_content)
    
    response = llm.invoke([system_msg, user_msg])
    return response.content


def refine_section_content(content: str, instruction: str, context: Optional[str] = None) -> str:
    """
    Refines existing section content based on natural language instructions (e.g. translate, rewrite, summarize).
    """
    check_gemini_config()
    
    llm = ChatGoogleGenerativeAI(
        model=settings.GEMINI_MODEL,
        google_api_key=settings.GEMINI_API_KEY,
        temperature=0.3  # Lower temperature for precise edits/refinement
    )
    
    system_msg = SystemMessage(content=(
        "You are SatoriAI, an expert AI document editor. "
        "Your task is to edit, refine, or rewrite the provided section content based on the user's instructions. "
        "Retain the markdown formatting (lists, bold/italic, tables, code blocks). "
        "Return ONLY the updated content. Do not include introductory notes, conversational text, "
        "or markdown wrapper tags like ```markdown ... ```. Output raw markdown content directly."
    ))
    
    user_content = f"Instruction: {instruction}\n\n"
    if context:
        user_content += f"Overall document context: {context}\n\n"
    user_content += f"Existing Section Content:\n---\n{content}\n---\n"
    
    user_msg = HumanMessage(content=user_content)
    
    response = llm.invoke([system_msg, user_msg])
    return response.content


def create_document_from_outline(db: Session, document_id: UUID, outline: DocumentOutline) -> Document:
    """
    Overwrites or populates a document's modules and sections based on a generated DocumentOutline.
    """
    doc = db.query(Document).filter(Document.id == document_id).first()
    if not doc:
        raise ValueError("Document not found")
        
    if outline.description and not doc.description:
        doc.description = outline.description
        
    # Delete existing modules/sections to overwrite the outline cleanly
    db.query(Module).filter(Module.document_id == document_id).delete()
    
    # Save the outline structure
    for m_idx, m_outline in enumerate(outline.modules):
        db_module = Module(
            document_id=document_id,
            title=m_outline.title,
            description=m_outline.description,
            order=m_idx
        )
        db.add(db_module)
        db.flush()  # Flush to get the database-generated ID
        
        def save_sections_recursive(sections_list: List[SectionOutline], parent_id: Optional[UUID] = None):
            for s_idx, s_outline in enumerate(sections_list):
                db_section = Section(
                    module_id=db_module.id,
                    parent_id=parent_id,
                    title=s_outline.title,
                    level=s_outline.level,
                    content="",  # Empty initially
                    order=s_idx
                )
                db.add(db_section)
                db.flush()
                
                if s_outline.children:
                    save_sections_recursive(s_outline.children, db_section.id)
                    
        save_sections_recursive(m_outline.sections)
        
    db.commit()
    db.refresh(doc)
    return doc


def generate_module_sections(module_title: str, module_desc: Optional[str] = None) -> List[SectionOutline]:
    """
    AI-generates a list of section outlines for a specific module title and description.
    """
    check_gemini_config()
    
    from pydantic import BaseModel, Field
    
    class ModuleSectionList(BaseModel):
        sections: List[SectionOutline] = Field(..., description="List of section outlines for this module")
        
    llm = ChatGoogleGenerativeAI(
        model=settings.GEMINI_MODEL,
        google_api_key=settings.GEMINI_API_KEY,
        temperature=0.7
    )
    
    structured_llm = llm.with_structured_output(ModuleSectionList)
    
    prompt = f"Create a comprehensive list of section headings (H1 to H5 hierarchy) for a document module titled '{module_title}'."
    if module_desc:
        prompt += f"\nModule Description/Context: {module_desc}"
        
    res = structured_llm.invoke(prompt)
    return res.sections


def run_ai_chat_agent(db: Session, user_id: UUID, message: str, document_id: Optional[UUID] = None) -> str:
    """
    Runs a tool-calling conversational agent that has access to the user's documents
    for searching, reading, and refining contents.
    """
    check_gemini_config()
    
    from langchain_core.tools import StructuredTool
    from langchain_core.messages import ToolMessage
    
    class DocumentTools:
        def __init__(self, db: Session, user_id: UUID):
            self.db = db
            self.user_id = user_id

        def list_documents(self, query: Optional[str] = None) -> str:
            """List the titles, descriptions, and IDs of your documents. Optionally filter by a search query."""
            q = self.db.query(Document).filter(Document.user_id == self.user_id)
            if query:
                q = q.filter(Document.title.ilike(f"%{query}%") | Document.description.ilike(f"%{query}%"))
            docs = q.all()
            if not docs:
                return "No documents found."
            return "\n".join([f"- Title: '{d.title}'\n  ID: {d.id}\n  Description: {d.description or 'None'}" for d in docs])

        def get_document_outline(self, document_id: str) -> str:
            """Get the module and section hierarchy outline of a specific document by its ID."""
            try:
                doc_uuid = UUID(document_id)
            except ValueError:
                return "Error: Invalid document ID UUID format."
            doc = self.db.query(Document).filter(Document.id == doc_uuid, Document.user_id == self.user_id).first()
            if not doc:
                return "Error: Document not found."
            lines = [f"Document: {doc.title} (ID: {doc.id})", f"Description: {doc.description or ''}\nOutline:"]
            for m in sorted(doc.modules, key=lambda x: x.order):
                lines.append(f"- Module: {m.title} (ID: {m.id})")
                top_sections = [s for s in m.sections if s.parent_id is None]
                top_sections.sort(key=lambda x: x.order)
                def add_sec(sec):
                    indent = "  " * (sec.level)
                    lines.append(f"{indent}- Section: '{sec.title}' (ID: {sec.id}, Level: H{sec.level})")
                    for c in sorted(sec.children, key=lambda x: x.order):
                        add_sec(c)
                for s in top_sections:
                    add_sec(s)
            return "\n".join(lines)

        def read_section_content(self, section_id: str) -> str:
            """Read the detailed text/markdown content of a specific section by its ID."""
            try:
                sec_uuid = UUID(section_id)
            except ValueError:
                return "Error: Invalid section ID UUID format."
            sec = self.db.query(Section).join(Module).join(Document).filter(
                Section.id == sec_uuid,
                Document.user_id == self.user_id
            ).first()
            if not sec:
                return "Error: Section not found."
            return f"Section Title: {sec.title}\nContent:\n{sec.content or 'Empty'}"

        def edit_section_content(self, section_id: str, instruction: str) -> str:
            """Edit, refine, or rewrite the content of a specific section by its ID using natural language instructions."""
            try:
                sec_uuid = UUID(section_id)
            except ValueError:
                return "Error: Invalid section ID UUID format."
            sec = self.db.query(Section).join(Module).join(Document).filter(
                Section.id == sec_uuid,
                Document.user_id == self.user_id
            ).first()
            if not sec:
                return "Error: Section not found."
            
            context = f"Document: {sec.module.document.title}."
            try:
                refined = refine_section_content(sec.content or "", instruction, context)
                sec.content = refined
                self.db.commit()
                return f"Section '{sec.title}' edited successfully. New Content:\n{refined}"
            except Exception as e:
                return f"Error: Failed to edit section: {str(e)}"
                
    tools = DocumentTools(db, user_id)
    langchain_tools = [
        StructuredTool.from_function(
            func=tools.list_documents,
            name="list_documents",
            description="List the titles, descriptions, and IDs of your documents. Optionally filter by a search query."
        ),
        StructuredTool.from_function(
            func=tools.get_document_outline,
            name="get_document_outline",
            description="Get the outline structure (modules and sections hierarchy) of a specific document by its ID."
        ),
        StructuredTool.from_function(
            func=tools.read_section_content,
            name="read_section_content",
            description="Read the detailed text/markdown content of a specific section by its ID."
        ),
        StructuredTool.from_function(
            func=tools.edit_section_content,
            name="edit_section_content",
            description="Edit or refine the content of a specific section by its ID using natural language instructions."
        )
    ]
    
    llm = ChatGoogleGenerativeAI(
        model=settings.GEMINI_MODEL,
        google_api_key=settings.GEMINI_API_KEY,
        temperature=0.4
    )
    llm_with_tools = llm.bind_tools(langchain_tools)
    
    prompt_context = (
        "You are SatoriAI, an expert document assistant. You help the user manage, search, read, "
        "and edit their documents. You have access to tools to view and modify the user's "
        "documents. Only search or modify documents belonging to the user. "
        "Explain what changes you have made clearly."
    )
    if document_id:
        prompt_context += f"\nThe user is currently focusing on document ID: {document_id}."
        
    messages = [
        SystemMessage(content=prompt_context),
        HumanMessage(content=message)
    ]
    
    for _ in range(5):
        res = llm_with_tools.invoke(messages)
        messages.append(res)
        
        if not res.tool_calls:
            break
            
        for tc in res.tool_calls:
            tool_name = tc["name"]
            tool_args = tc["args"]
            tool_id = tc["id"]
            
            if tool_name == "list_documents":
                output = tools.list_documents(**tool_args)
            elif tool_name == "get_document_outline":
                output = tools.get_document_outline(**tool_args)
            elif tool_name == "read_section_content":
                output = tools.read_section_content(**tool_args)
            elif tool_name == "edit_section_content":
                output = tools.edit_section_content(**tool_args)
            else:
                output = f"Tool '{tool_name}' not found."
                
            messages.append(ToolMessage(content=output, tool_call_id=tool_id))
            
    return res.content



